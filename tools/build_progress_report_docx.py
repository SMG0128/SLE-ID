from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "SLE项目当前进度与双板联调验收报告_2026-08-13.docx"
BLUE, DARK, NAVY = "2E74B5", "1F4D78", "20364C"
LIGHT, GRAY, GREEN, AMBER = "E8EEF5", "F2F4F7", "E2F0D9", "FFF2CC"
MUTED, INK = "667085", "202124"

def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr(); node = pr.find(qn("w:shd"))
    if node is None: node = OxmlElement("w:shd"); pr.append(node)
    node.set(qn("w:fill"), fill)

def margins(cell, top=80, start=120, bottom=80, end=120):
    pr = cell._tc.get_or_add_tcPr(); mar = pr.first_child_found_in("w:tcMar")
    if mar is None: mar = OxmlElement("w:tcMar"); pr.append(mar)
    for side, val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        n = mar.find(qn(f"w:{side}"))
        if n is None: n = OxmlElement(f"w:{side}"); mar.append(n)
        n.set(qn("w:w"), str(val)); n.set(qn("w:type"), "dxa")

def geometry(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    pr = table._tbl.tblPr; tw = pr.find(qn("w:tblW"))
    if tw is None: tw = OxmlElement("w:tblW"); pr.append(tw)
    tw.set(qn("w:w"), str(sum(widths))); tw.set(qn("w:type"), "dxa")
    ind = pr.find(qn("w:tblInd"))
    if ind is None: ind = OxmlElement("w:tblInd"); pr.append(ind)
    ind.set(qn("w:w"), "120"); ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcw is None: tcw = OxmlElement("w:tcW"); cell._tc.get_or_add_tcPr().append(tcw)
            tcw.set(qn("w:w"), str(widths[i])); tcw.set(qn("w:type"), "dxa")
            margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def font(run, size=10.5, bold=False, color=INK):
    run.font.name = "Microsoft YaHei"; rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei"); rpr.rFonts.set(qn("w:ascii"), "Calibri"); rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor.from_string(color)

def configure(doc):
    s = doc.sections[0]; s.page_width=Inches(8.5); s.page_height=Inches(11)
    s.top_margin=Inches(.58); s.bottom_margin=Inches(.58); s.left_margin=Inches(1); s.right_margin=Inches(1)
    s.header_distance=Inches(.4); s.footer_distance=Inches(.4)
    st=doc.styles["Normal"]; st.font.name="Microsoft YaHei"; st._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); st.font.size=Pt(10.5); st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.12
    for name,size,color,before,after in (("Heading 1",16,BLUE,9,5),("Heading 2",13,BLUE,7,4),("Heading 3",11.5,DARK,6,3)):
        st=doc.styles[name]; st.font.name="Microsoft YaHei"; st._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    p=s.header.paragraphs[0]; font(p.add_run("SLE-ID 硬件端与管理端 / 项目状态简报"),9,True,MUTED)
    p=s.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(p.add_run("项目进度报告 · 2026-08-13"),9,False,MUTED)

def para(doc,text,color=INK,after=6):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); font(p.add_run(text),10.5,False,color); return p

def callout(doc,label,text,fill=GREEN):
    t=doc.add_table(rows=1,cols=1); geometry(t,[9360]); c=t.cell(0,0); shade(c,fill); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(label+"  "),11,True,NAVY); font(p.add_run(text),10.5,False,NAVY); doc.add_paragraph().paragraph_format.space_after=Pt(1)

def table(doc,headers,rows,widths,status=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; trpr=t.rows[0]._tr.get_or_add_trPr(); head=OxmlElement("w:tblHeader"); trpr.append(head)
    for i,v in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,LIGHT); p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); font(p.add_run(v),9.2,True,NAVY)
    for values in rows:
        cells=t.add_row().cells
        for i,v in enumerate(values):
            p=cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i==status or len(v)<=10 else WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_after=Pt(0); font(p.add_run(v),8.9)
            if i==status: shade(cells[i],GREEN if v in ("通过","完成") else AMBER)
    geometry(t,widths); doc.add_paragraph().paragraph_format.space_after=Pt(1)

def bullet(doc,text):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Inches(.375); p.paragraph_format.first_line_indent=Inches(-.188); p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.10; font(p.add_run(text),10.2)

def build():
    doc=Document(); configure(doc)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(22); p.paragraph_format.space_after=Pt(4); font(p.add_run("SLE 项目当前进度与双板联调验收报告"),23,True,NAVY)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(18); font(p.add_run("Detector A / Detector B / 管理端 · 截至 2026 年 8 月 13 日"),12,False,MUTED)
    callout(doc,"总体结论","两块 BearPi H3863 已完成 A、B 实体联调并接入管理端。基础策略、事件、报警和断线恢复闭环均通过；真实授权执行仍需第三块 SLE 板运行 Card C。")
    doc.add_heading("1. 当前版本基线",1)
    table(doc,["项目","当前值"],[["通信协议","SLE Protocol V2"],["Detector A 固件","detector_a_h3863_all.fwpkg"],["A SHA256","5D4B54D56A06E5117154A446EEDB216C36DF4AD8524D170EF32BEFF59341F949"],["Detector B 固件","detector_b_h3863_all.fwpkg"],["B SHA256","2C49CB51DBAC604295CF4FAB75DDE4A7858EC330A25A7F86A17D5F1275E0B3E3"],["管理端后端","Node.js + TypeScript + SQLite，127.0.0.1:8080"],["实体硬件","2 × BearPi H3863（A、B）"]],[2160,7200])
    doc.add_heading("2. 本轮实体联调验收",1)
    table(doc,["编号","验收项","实际结果","结论"],[["T01","B 网关识别与双向心跳","首页显示“心跳正常”；设备心跳显示 0s","通过"],["T02","基础策略下发","连续下发两次成功，硬件策略版本递增","通过"],["T03","A→B 无线事件","A 收到 decision，B 正常返回未授权决策","通过"],["T04","B→后端→管理端事件","事件中心出现对应事件记录","通过"],["T05","未授权报警","策略设为报警后，报警中心出现记录","通过"],["T06","报警处理闭环","报警状态成功更新，首页统计同步","通过"],["T07","B 断线与恢复","USB 拔插后 COM、心跳及链路恢复","通过"],["T08","A 重启恢复","重新 bridge ready，新事件正常上报且未误去重","通过"],["T09","无 Card 安全拒绝","reason=1、exec=0，GPIO10 未误动作","通过"]],[620,2350,5110,1280],3)
    doc.add_heading("3. 已完成功能",1)
    for heading,items in (("3.1 Detector A",("SLE 客户端连接、服务发现与 bridge ready 状态管理。","区域进入、反向离开、超时等状态机测试入口。","Protocol V2 事件发送、ACK、超时重试和重启后唯一事件键。","认证中继核心已完成；最终授权只接受 B 的权威结果。")),("3.2 Detector B / 网关",("SLE 服务端、事件去重、决策回传和 GPIO10 执行器控制。","单 USB 串口网关：心跳、事件、报警、确认、策略与命令结果。","Card HMAC 认证权威核心、会话超时、计数器与防重放逻辑。","管理端离线识别、RAM 补传队列及命令重试。")),("3.3 管理端",("SQLite 数据库存储许可、设备、事件、报警、确认和审计数据。","REST + WebSocket 实时链路，能够识别 Detector B 合法心跳。","许可创建与基础策略下发、事件展示、未授权报警和报警处理闭环。","串口配置、断线重连、事务后 ACK、幂等与命令重试。"))):
        doc.add_heading(heading,2)
        for item in items: bullet(doc,item)
    doc.add_heading("4. 质量与修复状态",1)
    callout(doc,"关键修复","A/B 发布配置已明确禁用独立 Card C 组件；发布脚本会扫描最终链接符号，发现 Card 入口即停止打包。修复后 A/B 启动日志不再出现任何 [C] 行。",LIGHT)
    para(doc,"主机自动化回归已覆盖 A/B Protocol V2、Card 存储与服务、HMAC/防重放、B 认证权威、A 中继、B 单 USB 网关、Card 串口 DryRun、安全冒烟和管理端映射；本轮构建前后均通过。")
    doc.add_heading("5. 当前限制与待验证项",1)
    table(doc,["项目","当前状态","影响 / 处理方式"],[["Card C 实体板","待第三块板","无法进行真实 HMAC 授权、凭证写入及掉电保存验收"],["GPIO10 授权执行","待 Card C","安全版禁止用 auth ok 绕过，现阶段只能验证正确拒绝"],["管理员二次确认","待 Card C","需先建立真实认证会话，再验证确认后执行"],["真实 Channel Sounding","未接入","A 当前通过 demo enter 模拟检测输入"],["B 离线队列持久化","部分完成","当前以 RAM 补传为主，长期掉电需 Flash 队列"],["十项组合策略","基础子集完成","范围、时间、次数、方向等需扩展正式协议"],["后端真正重启","本轮未执行","重复启动仅证明原 8080 实例仍运行；后续补测配置与数据恢复"],["心跳页面展示","可用但不直观","取整且页面非实时刷新，持续显示 0s 属正常现象"]],[2450,1750,5160])
    doc.add_heading("6. 下一阶段建议",1)
    doc.add_heading("阶段 A：第三块板到手前",2)
    for x in ("将设备心跳改为自动刷新，并显示“刚刚 / 最后心跳时间”。","提供管理端一键启动脚本与清晰的进程/端口占用提示。","完成后端真正停止、重启、COM 配置恢复和 SQLite 数据持久化验收。","整理并提交源码、构建脚本、测试指南、固件和 SHA256SUMS.txt。"): bullet(doc,x)
    doc.add_heading("阶段 B：Card C 实体联调",2)
    for x in ("烧录 Card C 个性化固件，写入设备 ID、密钥与许可凭证。","验证服务发现、HMAC-SHA-256、会话计数器、防重放和掉电恢复。","验证授权成功后的 GPIO10 输出及管理员二次确认闭环。","完成后烧录 load_only 生产锁定包，关闭本地串口写卡入口。"): bullet(doc,x)
    doc.add_heading("7. 验收判定",1)
    callout(doc,"当前里程碑","A/B 双板 + 管理端基础链路可以判定为“已通过实体联调”。完整三端产品验收仍以 Card C 实体认证、GPIO 执行、二次确认和凭证持久化全部通过为完成条件。",AMBER)
    para(doc,"报告生成日期：2026-08-13。测试依据为当天实体板操作结果、当前项目源码、官方 SDK 构建产物及自动化测试输出。",MUTED,0)
    doc.core_properties.title="SLE 项目当前进度与双板联调验收报告"; doc.core_properties.subject="Detector A / Detector B / 管理端项目进度"; doc.core_properties.author="SLE-ID 项目组"
    OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); print(OUT)

if __name__ == "__main__": build()

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "硬件端与管理端兼容性问题报告_2026-08-09.docx"
BLUE = "2E74B5"
DARK = "1F4D78"
LIGHT = "F2F4F7"
RED = "C00000"
AMBER = "C65911"
GREEN = "548235"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        elem = tbl_cell_mar.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            tbl_cell_mar.append(elem)
        elem.set(qn("w:w"), str(value))
        elem.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_keep(paragraph, next_too=False):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepLines")
    p_pr.append(keep)
    if next_too:
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)


def add_run(paragraph, text, bold=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return run


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    add_run(p, text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    add_run(p, text)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, label in enumerate(headers):
        set_cell_shading(header.cells[idx], LIGHT)
        p = header.cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, label, bold=True, color=DARK)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_run(p, str(value))
    set_table_geometry(table, widths)
    before = doc.paragraphs[-1] if doc.paragraphs else None
    if before:
        before.paragraph_format.space_after = Pt(4)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 24, DARK, 0, 8),
        ("Subtitle", 12, "666666", 0, 12),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name.startswith("Heading") or name == "Title"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, "SLE-ID · 硬件/管理端兼容性检查", color="808080", size=9)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "内部联调文档  |  ", color="808080", size=9)
    page = OxmlElement("w:fldSimple")
    page.set(qn("w:instr"), "PAGE")
    footer._p.append(page)


def add_status_box(doc, label, text, color):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_shading(table.cell(0, 0), color)
    p0 = table.cell(0, 0).paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p0, label, bold=True, color="FFFFFF")
    p1 = table.cell(0, 1).paragraphs[0]
    add_run(p1, text)
    set_table_geometry(table, [1500, 7860])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    add_run(title, "硬件端与管理端兼容性问题报告", bold=True, color=DARK, size=24)
    subtitle = doc.add_paragraph(style="Subtitle")
    add_run(subtitle, "Detector A / Detector B / MilkWa-StarFollow 管理端", color="666666", size=12)
    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    meta_rows = (
        ("日期", "2026-08-09"),
        ("硬件工程", str(ROOT)),
        ("管理端快照", r"D:\naiwa1\SLE-ID-Admin-Backend\SLE-ID-Admin-Backend"),
    )
    for row, pair in zip(meta.rows, meta_rows):
        set_cell_shading(row.cells[0], LIGHT)
        add_run(row.cells[0].paragraphs[0], pair[0], bold=True, color=DARK)
        add_run(row.cells[1].paragraphs[0], pair[1])
    set_table_geometry(meta, [1700, 7660])
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(10)
    p_bdr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    p_bdr.append(borders)

    doc.add_heading("执行摘要", level=1)
    add_status_box(doc, "结论", "协议与字段映射已兼容；管理端前端可构建，但尚无可运行后端，因此真实串口→数据库→WebSocket→页面的端到端链路当前无法验收。", BLUE)
    add_status_box(doc, "修改原则", "本轮仅修改硬件工程与兼容工具；管理端源码零改动。管理端目录只因构建测试生成 node_modules/ 和 dist/。", GREEN)
    p = doc.add_paragraph()
    add_run(p, "已完成：", bold=True)
    add_run(p, "B 端单 USB Protocol V2 网关、1 秒心跳、策略/事件/告警/确认/命令结果、ACK 重试与 32 条 RAM 补传队列；管理 DTO 映射工具与自动测试。")
    p = doc.add_paragraph()
    add_run(p, "仍阻塞：", bold=True, color=RED)
    add_run(p, "管理端 server/ 只有 README，无串口驱动、SQLite、REST 和 WebSocket 源码。")

    doc.add_heading("1. 检查范围与方法", level=1)
    add_bullet(doc, "静态核查管理端目录、前端类型、API 层、WebSocket 占位与 server/ 内容。")
    add_bullet(doc, "使用 npm 官方源完成依赖安装，并执行 vue-tsc --noEmit 与 Vite 生产构建。")
    add_bullet(doc, "对照管理端 Device、EventLogItem、AlarmType、LicensePolicy 字段，核查硬件 Protocol V2。")
    add_bullet(doc, "新增离线映射 SelfTest、网关边界单测、整套主机回归，并使用官方 SDK 全量编译 A/B。")
    p = doc.add_paragraph()
    add_run(p, "注意：", bold=True, color=AMBER)
    add_run(p, "因为管理端缺少真实后端，本报告中的“兼容通过”指消息协议、字段语义和参考驱动自测通过，不等同于真实整机端到端通过。")

    doc.add_heading("2. 兼容性矩阵", level=1)
    add_table(doc,
              ["管理能力", "硬件消息/机制", "当前状态", "管理端待办"],
              [
                  ("设备在线、固件、策略版本、运行时长", "HEARTBEAT 0x30，1 秒周期", "硬件完成", "入库并以超时判定在线"),
                  ("事件日志", "EVENT_REPORT 0x62", "硬件完成", "复合唯一键并 WebSocket 推送"),
                  ("告警中心", "ALERT_REPORT 0x63", "硬件完成", "原因码映射 8 类 AlarmType"),
                  ("强制确认", "CONFIRM_REQUEST/RESULT 0x64/0x65", "硬件完成", "补确认 UI/API，并复用 requestId+eventId"),
                  ("策略发布与结果", "POLICY_SYNC/RESULT 0x60/0x61", "基础子集完成", "只在收到结果后标记成功"),
                  ("命令幂等", "requestId + COMMAND_RESULT 0x66", "硬件完成", "持久化请求状态"),
                  ("断线补传", "32 条 RAM 队列 + ACK/重试", "阶段完成", "主机 ACK；硬件后续升级 Flash"),
                  ("实时生命周期", "只有最终事件/确认", "不完整", "定义中间状态协议"),
              ], [2150, 2500, 1450, 3260])

    doc.add_heading("3. 问题清单与处理结果", level=1)
    add_table(doc,
              ["级别", "问题", "影响", "处理/责任"],
              [
                  ("P0", "管理端没有可运行后端", "无法真实连接 B、落库和推送页面", "未改管理端；管理端队友实现 server"),
                  ("P1", "Host 策略 requestId 与 B 确认 requestId 数字碰撞", "合法请求可能被误判冲突", "已在硬件按 requestId+commandType 分域"),
                  ("P1", "过期确认仍可能留在可靠队列", "超时后主机仍收到无效确认", "已在终态报告时删除旧确认帧"),
                  ("P1", "队列满时确认未入队却被标 active", "状态机等待一个无法送达的请求", "已仅在成功入队后激活"),
                  ("P1", "10 项管理策略只能映射基础子集", "页面显示成功会造成能力误报", "文档约束返回部分支持；协议后续扩展"),
                  ("P1", "离线记录仅 RAM 32 条", "B 重启会丢待传记录", "保留计数；后续实现 Flash 256 条"),
                  ("P2", "文本日志与二进制帧共用 UART", "增加解析、带宽与交错风险", "兼容工具支持重同步；发布版宜限流/关闭文本"),
                  ("P2", "管理端前端构建 chunk >500 kB", "首屏体积偏大", "不影响兼容；前端后续拆包"),
                  ("P2", "npm 报告 6 项依赖漏洞", "供应链维护风险", "未强制升级；前端队友评估后处理"),
              ], [700, 2800, 2600, 3260])

    doc.add_heading("4. 关键字段与唯一性约定", level=1)
    doc.add_heading("4.1 设备心跳", level=2)
    add_bullet(doc, "sourceId → Device.id；firmwareVersion 0x00010002 → v1.0.2；policyVersion 2 → POL-2。")
    add_bullet(doc, "uptimeMs → uptimeSeconds；queueDepth、queueOverflows、framesSent 进入设备诊断字段。")
    add_bullet(doc, "设备名称和位置由管理端 SQLite 按 sourceId 关联，禁止写死在固件。")
    doc.add_heading("4.2 事件和时间", level=2)
    p = doc.add_paragraph()
    add_run(p, "事件唯一键：", bold=True)
    add_run(p, "EV-{sourceId}-{bootId}-{eventId}", bold=True, color=DARK)
    add_run(p, "。A 重启后 eventId 从 1 开始，也不会覆盖历史。")
    add_bullet(doc, "cardAnonId 映射匿名 CARD-XXXXXXXX；持有人姓名只由管理数据库关联，硬件不上传姓名。")
    add_bullet(doc, "硬件时间为启动后毫秒数；管理端必须同时记录主机接收时间，不能将其直接当作 UTC。")
    doc.add_heading("4.3 告警映射", level=2)
    add_table(doc, ["硬件原因", "管理 AlarmType"], [
        ("NO_PERMISSION / OUT_OF_SCOPE", "unauthorized"),
        ("EXPIRED", "license_expired"),
        ("LOST", "lost_report"),
        ("KEY_FAILED", "key_failed"),
        ("REPLAY_SUSPECTED", "suspected_replay"),
        ("CONFIRM_REJECTED / TIMEOUT / OFFLINE", "confirm_rejected"),
        ("EXECUTION_FAILED", "execute_failed"),
        ("未登记 sourceId（后端产生）", "unknown_device"),
    ], [4700, 4660])

    doc.add_heading("5. 策略兼容边界", level=1)
    p = doc.add_paragraph()
    add_run(p, "可直接执行：", bold=True, color=GREEN)
    add_run(p, "allowExecute、forceConfirm/用户确认、offlineAllowed、unauthorizedAction=alarm。")
    p = doc.add_paragraph()
    add_run(p, "暂不能宣称完整支持：", bold=True, color=AMBER)
    add_run(p, "recordEvent=false、scope、timeLimit、usageLimit、direction、huaweiFallback。")
    add_bullet(doc, "B 为安全审计始终上报事件，因此 recordEvent=false 当前不支持。")
    add_bullet(doc, "scope/time/usage/direction 在认证或事件模型中有部分基础，但未形成完整下发与持久执行闭环。")
    add_bullet(doc, "huaweiFallback 属于 App/后端临时凭证流程，不应塞入 B 本地布尔策略。")
    add_bullet(doc, "管理后端应返回“支持/部分支持/不支持”，不能只显示“策略更新成功”。")

    doc.add_heading("6. 本轮硬件修改", level=1)
    add_number(doc, "新增 Detector B 单 USB Protocol V2 网关：心跳、策略、事件、告警、确认、命令结果与可靠发送。")
    add_number(doc, "心跳周期从 2 秒统一为 1 秒，满足管理端在线监测粒度。")
    add_number(doc, "修复 requestId 命名空间碰撞、过期确认残留、满队列 active 状态三个边界缺陷。")
    add_number(doc, "新增 tools/admin_gateway_compat.ps1：混合流解析、CRC、自动 ACK、Host 心跳和 DTO 风格 NDJSON。")
    add_number(doc, "新增映射 SelfTest 和 3 组网关回归测试，并纳入 tests/run_tests.ps1。")
    add_number(doc, "更新协议、兼容说明、README、工作进度和修改日志。")

    doc.add_heading("7. 验证结果", level=1)
    add_table(doc, ["项目", "结果", "说明"], [
        ("管理端 TypeScript + Vite 生产构建", "通过", "2257 modules；仅 chunk/PURE 警告"),
        ("管理字段映射 SelfTest", "通过", "Device/Event/Alarm/Confirm；1 秒心跳"),
        ("硬件主机回归", "通过", "协议、A/B、Card、认证、中继、网关、工具"),
        ("Detector A 官方 SDK 全量构建", "通过", "编译、签名、打包"),
        ("Detector B 官方 SDK 全量构建", "通过", "编译、签名、打包"),
        ("真实串口→SQLite→WebSocket→页面", "未执行", "管理端真实后端不存在"),
    ], [3100, 1300, 4960])

    doc.add_heading("8. 发布产物与校验值", level=1)
    add_table(doc, ["文件", "SHA256"], [
        ("detector_a_h3863_all.fwpkg", "0C80BFC5F86842624DA666EDE70204EEA486601F6FC8A6F0134350C553966291"),
        ("detector_b_h3863_all.fwpkg", "DF58FA33A8F1DF826428F6A92508E84555E6C428D32BB89AD1FFF316C654F5F9"),
        ("card_c_h3863_ram_all.fwpkg", "4F5D6364ED4B7C206D7BAA45098143724FA004775B65F3DA9CB4B0094451BB56"),
        ("SLE_AB_H3863_source.zip", "859C414A94EF32203F9F2A52453E88335360AC1E97A0DA465E60A690F0FF5C69"),
    ], [3300, 6060])

    doc.add_heading("9. 管理端后端最小实施清单", level=1)
    add_number(doc, "以 115200/8N1 打开 B 串口，按魔数 53 4C 流式重同步；不能按文本行读取。")
    add_number(doc, "每秒发送 Host HEARTBEAT，对 ACK_REQUIRED 消息立即回复 ACK。")
    add_number(doc, "将心跳、事件和告警写入 SQLite，再通过 WebSocket 推送前端。")
    add_number(doc, "REST 保持 {code, message, data}；策略收到 POLICY_RESULT 后才标记发布成功。")
    add_number(doc, "确认结果携带原 requestId+eventId；超时确认不得重试成批准。")
    add_number(doc, "数据库对 sourceId+bootId+eventId 建唯一索引；接收时间与设备单调时间分开保存。")

    doc.add_heading("10. 下一轮验收步骤", level=1)
    add_number(doc, "管理端队友提交可运行 server 源码和启动说明。")
    add_number(doc, "先用 tools/admin_gateway_compat.ps1 -Action Listen -Port COMx 验证 B 帧、CRC、ACK 和 NDJSON。")
    add_number(doc, "后端接入同一协议后，验证设备在线、事件入库、告警推送、策略回执和确认闭环。")
    add_number(doc, "断开/恢复 USB，验证 ACK 重发和补传；重启 B，明确 RAM 队列丢失属于当前已知限制。")
    add_number(doc, "对管理端 10 项策略逐项显示支持等级，禁止未实现项返回成功。")

    doc.add_heading("附录：兼容工具命令", level=1)
    commands = [
        r"powershell -NoProfile -ExecutionPolicy Bypass -File .\tools\admin_gateway_compat.ps1 -Action SelfTest",
        r"powershell -NoProfile -ExecutionPolicy Bypass -File .\tools\admin_gateway_compat.ps1 -Action Listen -Port COM6",
        r"powershell -NoProfile -ExecutionPolicy Bypass -File .\tools\admin_gateway_compat.ps1 -Action Policy -Port COM6 -PermissionId 7 -OrganizationId 100 -PolicyVersion 2 -AllowExecute -OfflineAllowed -AlertOnDenial",
        r"powershell -NoProfile -ExecutionPolicy Bypass -File .\tools\admin_gateway_compat.ps1 -Action Confirm -Port COM6 -RequestId 77 -EventId 9 -ConfirmResult Approve",
    ]
    for command in commands:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.1)
        set_cell = None
        run = add_run(p, command, color=DARK, size=9)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        set_keep(p)

    doc.core_properties.title = "硬件端与管理端兼容性问题报告"
    doc.core_properties.subject = "SLE Detector A/B 与 MilkWa-StarFollow 管理端兼容性"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "SLE, H3863, 管理端, 兼容性, Protocol V2"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

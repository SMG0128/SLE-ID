import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "ADMIN_BACKEND_DESIGN_V1.md"
OUTPUT = ROOT / "outputs" / "SLE-ID管理端后端与硬件通信方案_审阅版V1_2026-08-09.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "548235"
GOLD = "7A5A00"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_left_border(paragraph, color=BLUE, size=18, space=8):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    borders.append(left)


def set_paragraph_bottom_border(paragraph, color=BLUE, size=12):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def set_keep(paragraph, with_next=False):
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:keepLines"))
    if with_next:
        p_pr.append(OxmlElement("w:keepNext"))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def apply_table_geometry(table, widths):
    if sum(widths) != 9360:
        raise ValueError(f"Table widths must total 9360 DXA, got {sum(widths)}")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        set_row_cant_split(row)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_numbering(doc, fmt, text):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    if fmt == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def add_inline(paragraph, text, size=None, color=None):
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(text[cursor:match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=(size or 10), color=DARK_BLUE)
        cursor = match.end()
    if cursor < len(text):
        set_run_font(paragraph.add_run(text[cursor:]), size=size, color=color)


def add_callout(doc, label, text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.10
    set_paragraph_shading(p, CALLOUT)
    set_paragraph_left_border(p, color=color)
    set_run_font(p.add_run(label + "  "), size=11, color=color, bold=True)
    add_inline(p, text, size=11, color=INK)
    set_keep(p)
    return p


def choose_widths(headers):
    count = len(headers)
    joined = " ".join(headers)
    if count == 2:
        if "SHA" in joined:
            return [3300, 6060]
        if "code" in joined or "含义" in joined:
            return [1600, 7760]
        return [2600, 6760]
    if count == 3:
        return [2300, 2500, 4560]
    if count == 4:
        if "offset" in joined or "field" in joined:
            return [900, 800, 2600, 5060]
        if "方向" in joined and "类型" in joined:
            return [1300, 1500, 3000, 3560]
        return [1500, 2600, 2300, 2960]
    return [9360 // count] * (count - 1) + [9360 - (9360 // count) * (count - 1)]


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_row = table.rows[0]
    set_repeat_header(header_row)
    for index, header in enumerate(headers):
        set_cell_shading(header_row.cells[index], LIGHT_GRAY)
        p = header_row.cells[index].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_inline(p, header, size=9.5, color=DARK_BLUE)
        for run in p.runs:
            run.bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if len(value) < 18 and index == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, value.replace("<br>", "\n"), size=9.5)
    apply_table_geometry(table, choose_widths(headers))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)
    return table


def add_code_block(doc, lines, language):
    if language == "mermaid":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.right_indent = Inches(0.12)
        set_paragraph_shading(p, "E8EEF5")
        set_paragraph_left_border(p, color=BLUE, size=12)
        flow = "Card C  ⇄ SLE ⇄  Detector A  ⇄ SLE ⇄  Detector B  ⇄ USB/UART ⇄  本机后端  ⇄ REST/WS ⇄  Vue 管理端"
        set_run_font(p.add_run(flow), size=10.5, color=INK, bold=True)
        set_keep(p)
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    set_paragraph_shading(p, LIGHT_GRAY)
    set_paragraph_left_border(p, color="98A2B3", size=8)
    run = p.add_run("\n".join(lines))
    set_run_font(run, name="Consolas", size=8.5, color=INK)
    set_keep(p)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run("SLE-ID  |  管理端后端与硬件通信方案  |  审阅版 V1"), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run("内部技术审阅  |  第 "), size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    set_run_font(footer.add_run(" 页"), size=8.5, color=MUTED)


def add_masthead(doc):
    title = doc.add_paragraph(style="Title")
    set_run_font(title.add_run("SLE-ID 管理端后端与硬件通信方案"), size=24, color=INK, bold=True)
    subtitle = doc.add_paragraph(style="Subtitle")
    set_run_font(subtitle.add_run("基于现有 Vue 3 前端与硬件 Protocol V2 的审阅版设计"), size=13, color=MUTED)
    metadata = (
        ("版本", "审阅版 V1"),
        ("日期", "2026-08-09"),
        ("范围", "本机后端、SQLite、REST、WebSocket、单 USB 串口与 Detector B"),
        ("状态", "待项目成员审阅；尚未开始实现管理端后端"),
    )
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        set_run_font(p.add_run(label + "："), size=10.5, color=INK, bold=True)
        set_run_font(p.add_run(value), size=10.5, color=INK)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(10)
    set_paragraph_bottom_border(rule, BLUE, 12)
    add_callout(
        doc,
        "建议定案",
        "Node.js/TypeScript 本机模块化单体；电脑只连接 B 的单 USB；硬件继续使用 Protocol V2 二进制帧；REST 负责查询和操作，WebSocket 负责实时推送；事件先落 SQLite 再 ACK。",
        color=BLUE,
    )
    add_callout(
        doc,
        "审阅重点",
        "请重点确认技术栈、127.0.0.1:8080、B 单 USB、ACK 落盘语义、第一版单活动策略，以及未实现硬件能力不得伪造成功。",
        color=GOLD,
    )


def parse_markdown(doc, text, bullet_num_id, decimal_num_id):
    lines = text.splitlines()
    index = 0
    in_code = False
    code_language = ""
    code_lines = []
    before_first_h2 = True
    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        stripped = line.strip()
        if in_code:
            if stripped.startswith("```"):
                add_code_block(doc, code_lines, code_language)
                in_code = False
                code_language = ""
                code_lines = []
            else:
                code_lines.append(line)
            index += 1
            continue
        if stripped.startswith("```"):
            in_code = True
            code_language = stripped[3:].strip().lower()
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# "):
            index += 1
            continue
        if before_first_h2 and not stripped.startswith("## "):
            index += 1
            continue
        if stripped.startswith("## "):
            before_first_h2 = False
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, stripped[3:])
            set_keep(p, with_next=True)
            index += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, stripped[4:])
            set_keep(p, with_next=True)
            index += 1
            continue
        if stripped.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, stripped[5:])
            set_keep(p, with_next=True)
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[index + 1]):
            def parse_row(value):
                return [cell.strip() for cell in value.strip().strip("|").split("|")]
            headers = parse_row(line)
            index += 2
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                row = parse_row(lines[index])
                if len(row) < len(headers):
                    row.extend([""] * (len(headers) - len(row)))
                elif len(row) > len(headers):
                    row = row[:len(headers) - 1] + [" | ".join(row[len(headers) - 1:])]
                rows.append(row)
                index += 1
            add_table(doc, headers, rows)
            continue
        if re.match(r"^-\s+", stripped):
            p = doc.add_paragraph()
            apply_num(p, bullet_num_id)
            add_inline(p, re.sub(r"^-\s+", "", stripped))
            set_keep(p)
            index += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph()
            apply_num(p, decimal_num_id)
            add_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
            set_keep(p)
            index += 1
            continue
        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            nxt = lines[index].strip()
            if not nxt or nxt.startswith(("#", "```", "|", "- ")) or re.match(r"^\d+\.\s+", nxt):
                break
            paragraph_lines.append(nxt)
            index += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))
        set_keep(p)


def build():
    doc = Document()
    configure_styles(doc)
    bullet_num_id = add_numbering(doc, "bullet", "•")
    decimal_num_id = add_numbering(doc, "decimal", "%1.")
    add_masthead(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"), bullet_num_id, decimal_num_id)
    doc.core_properties.title = "SLE-ID 管理端后端与硬件通信方案（审阅版 V1）"
    doc.core_properties.subject = "本机后端、单 USB Protocol V2、REST、WebSocket 与 SQLite 设计"
    doc.core_properties.author = "SLE-ID 项目组"
    doc.core_properties.keywords = "SLE, H3863, 管理端, Node.js, SQLite, Protocol V2"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

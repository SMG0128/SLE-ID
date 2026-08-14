from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "SLE-ID移动端设计方案_审阅版V1_2026-08-14.docx"
ASSET = ROOT / "tmp" / "mobile_design_architecture.png"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
GREEN = "18794E"
AMBER = "8A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_font(run, size=None, bold=None, color=None, name="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Title", 25, NAVY, 0, 4),
        ("Subtitle", 13, MUTED, 0, 14),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        if style_name == "Title":
            ppr = style._element.get_or_add_pPr()
            border = ppr.find(qn("w:pBdr"))
            if border is not None:
                ppr.remove(border)


def add_numbering(doc, fmt, text, left=720, hanging=360):
    numbering = doc.part.numbering_part.element
    existing = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    ppr.append(ind)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rpr.append(rfonts)
    lvl.append(rpr)
    abstract.append(lvl)
    # OOXML requires every abstractNum to precede the concrete num records.
    # Inserting after existing num nodes makes Word repair the numbering part,
    # which can remap bullets to decimal numbering.
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)
    nums = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(nums, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def numbered_paragraph(doc, text, num_id, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    r = p.add_run(text)
    set_font(r, color=color)
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text="", bold_lead=None, color=INK, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = keep
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2, color=color)
    else:
        r = p.add_run(text)
        set_font(r, color=color)
    return p


def add_callout(doc, label, text, fill=CALLOUT, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label + "  ")
    set_font(r1, bold=True, color=accent)
    r2 = p.add_run(text)
    set_font(r2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_table(doc, headers, rows, widths, header_fill=LIGHT, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        set_cell_shading(hdr.cells[i], header_fill)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, size=font_size, bold=True, color=NAVY)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            if i == 0 and len(headers) <= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_font(r, size=font_size, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(lines)
    set_font(r, size=9, color="344054", name="Consolas", east_asia="Microsoft YaHei")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=MUTED)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("SLE-ID 移动端设计方案")
    set_font(r, size=9, color=MUTED)
    hp_pr = hp._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), "D7DBE2")
    border.append(bottom)
    hp_pr.append(border)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    r1 = fp.add_run("内部审阅 · 2026-08-14   |   第 ")
    set_font(r1, size=9, color=MUTED)
    add_page_number(fp)
    r2 = fp.add_run(" 页")
    set_font(r2, size=9, color=MUTED)


def make_architecture_image():
    ASSET.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1800, 980
    img = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    bold_path = Path("C:/Windows/Fonts/msyhbd.ttc")
    font = ImageFont.truetype(str(font_path), 32)
    small = ImageFont.truetype(str(font_path), 25)
    bold = ImageFont.truetype(str(bold_path if bold_path.exists() else font_path), 34)

    def box(x, y, w, h, title, sub, fill, outline):
        draw.rounded_rectangle((x, y, x + w, y + h), radius=24, fill=fill, outline=outline, width=4)
        tb = draw.textbbox((0, 0), title, font=bold)
        draw.text((x + (w - (tb[2] - tb[0])) / 2, y + 28), title, font=bold, fill="#17365D")
        sb = draw.multiline_textbbox((0, 0), sub, font=small, spacing=9, align="center")
        draw.multiline_text((x + (w - (sb[2] - sb[0])) / 2, y + 88), sub, font=small, fill="#344054", spacing=9, align="center")

    def arrow(x1, y1, x2, y2, label, color="#2E74B5"):
        draw.line((x1, y1, x2, y2), fill=color, width=6)
        if x1 == x2:
            draw.polygon([(x2, y2), (x2 - 13, y2 - 22), (x2 + 13, y2 - 22)], fill=color)
        else:
            draw.polygon([(x2, y2), (x2 - 22, y2 - 13), (x2 - 22, y2 + 13)], fill=color)
        if label:
            bb = draw.textbbox((0, 0), label, font=small)
            lx = (x1 + x2) / 2 - (bb[2] - bb[0]) / 2
            ly = (y1 + y2) / 2 - (bb[3] - bb[1]) / 2 - 5
            draw.rounded_rectangle((lx - 12, ly - 6, lx + (bb[2] - bb[0]) + 12, ly + (bb[3] - bb[1]) + 8), radius=8, fill="#FFFFFF")
            draw.text((lx, ly), label, font=small, fill=color)

    box(70, 70, 430, 175, "HarmonyOS 手机 App", "领卡 · 状态同步 · 二次确认\nSLE 写卡 · 用户安全设置", "#E8F2FF", "#2E74B5")
    box(685, 70, 430, 175, "Node.js 本地后端", "移动端 API · WebSocket\nSQLite · USB 网关", "#EEF8F3", "#18794E")
    box(1300, 70, 430, 175, "管理端 Ark Web", "许可发布 · 邀请码\n事件 · 报警 · 设备状态", "#F7F1FF", "#7A4FA3")
    arrow(500, 157, 685, 157, "HTTPS / WS")
    arrow(1115, 157, 1300, 157, "HTTP / WS")

    box(70, 410, 430, 175, "WS63 Card C", "匿名身份载体 · 最多8项许可\n认证 · 防重放 · 原子写入", "#FFF6E5", "#A86A00")
    box(685, 410, 430, 175, "Detector A", "目标发现 · CS/状态机\n完整经过判断 · Card认证", "#FFF1F1", "#B54747")
    box(1300, 410, 430, 175, "Detector B", "策略判断 · 二次确认等待\n执行器 · 报警 · USB网关", "#FFF1F1", "#B54747")
    arrow(500, 497, 685, 497, "SLE 认证")
    arrow(1115, 497, 1300, 497, "SLE V2")
    arrow(1515, 410, 1515, 275, "USB V2")
    arrow(285, 245, 285, 410, "SLE 写卡")

    draw.rounded_rectangle((70, 720, 1730, 900), radius=22, fill="#F7F8FA", outline="#D0D5DD", width=3)
    draw.text((105, 752), "链路边界", font=bold, fill="#17365D")
    notes = [
        "网络控制面：手机 ↔ 后端，传 JSON；负责领卡、同步、确认。",
        "写卡面：手机 ↔ Card C，传 Protocol V2 二进制帧；不经过 Detector B。",
        "检测面：Card C ↔ A ↔ B ↔ 后端；无确认策略时手机不在实时执行链路中。",
    ]
    for i, note in enumerate(notes):
        draw.ellipse((110, 812 + i * 32, 124, 826 + i * 32), fill="#2E74B5")
        draw.text((140, 798 + i * 32), note, font=small, fill="#344054")
    img.save(ASSET, quality=95)


def build_document():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    make_architecture_image()
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    bullet_id = add_numbering(doc, "bullet", "•")
    decimal_id = add_numbering(doc, "decimal", "%1.")

    # First-page memo masthead
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("技术方案 / 审阅版")
    set_font(r, size=10, bold=True, color=BLUE)
    title = doc.add_paragraph(style="Title")
    title.add_run("SLE-ID 移动端设计方案")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("基于现有硬件端、管理端与初版三端开发文档的落地设计")

    metadata = [
        ("文档版本", "V1.0（审阅版）"),
        ("发布日期", "2026年8月14日"),
        ("适用范围", "HarmonyOS ArkUI 手机端第一版本"),
        ("现有基础", "Detector A/B + Card C + Node.js/SQLite 管理端"),
        ("建议决策", "网络控制面与 SLE 写卡面分离的双通道架构"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(label + "：")
        set_font(r1, bold=True, color=NAVY)
        r2 = p.add_run(value)
        set_font(r2, color=INK)

    add_callout(
        doc,
        "建议冻结",
        "移动端通过 HTTP/WebSocket 与本地后端交换用户数据，通过 SLE 直接向 Card C 写入凭证；手机不直接连接 Detector B，SLE 链路使用二进制协议，网络接口使用 JSON。",
        fill=PALE_BLUE,
        accent=NAVY,
    )

    add_heading(doc, "1. 设计依据与目标", 1)
    add_para(doc, "本方案依据初版《SLE无感身份识别与事件确认系统——三端开发分工与开发路线》，并结合当前已经完成的 A/B 双板、Card C、USB 网关和管理端本地后端进行设计。移动端不是管理网页的简单封装，而是个人凭证管理中心、二次确认终端和 Card C 安全写卡端。")
    for text in [
        "复用现有 Detector A/B、Protocol V2、USB 网关、Node.js 后端和 SQLite 数据库。",
        "完成邀请码领卡、卡片状态、本地持久化、实时确认和真实 SLE 写卡。",
        "管理端强制二次确认不可被手机关闭；用户只能额外增加保护。",
        "任何写卡、确认和状态同步均以回执证明完成，禁止仅更新界面后宣称成功。",
    ]:
        numbered_paragraph(doc, text, bullet_id)

    add_heading(doc, "2. 总体架构", 1)
    add_para(doc, "第一版本采用“双通道、三链路”结构：网络控制面负责用户和管理数据，SLE 写卡面负责凭证落卡，检测面继续沿用现有 Card C—A—B—后端链路。")
    doc.add_picture(str(ASSET), width=Inches(6.45))
    cap = doc.paragraphs[-1]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("图1  移动端与现有三端系统的正式连接关系")
    set_font(r, size=9, color=MUTED)

    add_heading(doc, "2.1 三条业务链路", 2)
    add_table(
        doc,
        ["链路", "路径", "协议与职责"],
        [
            ("领卡与同步", "手机 ↔ 本地后端", "HTTPS/HTTP + JSON；邀请码、卡片状态、用户设置、挂失冻结。"),
            ("实体卡写入", "手机 ↔ Card C", "SLE Protocol V2 二进制帧；发现、认证、分片写入、提交、读回。"),
            ("二次确认", "B → 后端 → 手机 → 后端 → B", "USB V2 + WebSocket/REST；同意、拒绝、超时、离线。"),
            ("日常检测", "Card C ↔ A ↔ B ↔ 后端", "SLE + USB；无确认策略时手机不参与实时执行。"),
        ],
        [1800, 2800, 4760],
    )

    add_heading(doc, "3. 当前进度与差距", 1)
    add_table(
        doc,
        ["领域", "当前可复用能力", "移动端落地前必须补齐"],
        [
            ("Detector A/B", "事件状态机、认证转发、策略、确认等待、GPIO执行、Protocol V2。", "继续沿用；补充手机确认和真实 Card C 的三板验收。"),
            ("USB/后端", "串口心跳、事件、报警、策略下发、确认结果、SQLite、WebSocket。", "增加手机专用鉴权、用户数据隔离和局域网访问。"),
            ("Card C", "8项许可、78字节凭证、分片/CRC32、双槽提交、HMAC认证。", "当前远端 SLE 管理命令被拒绝；增加安全写卡票据并完成 NV 真板验收。"),
            ("手机工程", "已有部分 ArkUI 页面和模拟交互。", "恢复缺失的 models/stores/mock/页面文件，重新建立可编译工程。"),
            ("现有 API", "邀请码兑换、卡片、许可、待确认和确认决策。", "现有接口面向管理员且返回全局数据；不能直接暴露给普通手机用户。"),
        ],
        [1500, 3720, 4140],
        font_size=9,
    )
    add_callout(doc, "关键冲突", "初版文档要求手机通过 SLE 写 Card C，但当前 Card C 固件只允许本地串口 Host 写卡。移动端实施必须同时补充 Card C 的远端安全写卡授权，否则只能得到演示界面，无法完成真实闭环。", fill="FFF4E5", accent=AMBER)

    add_heading(doc, "4. 移动端应用结构", 1)
    add_heading(doc, "4.1 页面与用户流程", 2)
    add_table(
        doc,
        ["页面", "核心功能", "关键状态"],
        [
            ("卡包主页", "展示全部数字许可及实体卡绑定状态；下拉刷新。", "未写入、正常、即将到期、冻结、挂失、过期、撤销。"),
            ("邀请码领卡", "预览签发方、范围、有效期、次数和确认要求；确认兑换。", "预览中、已领取、邀请码无效/过期/已使用。"),
            ("SLE写卡向导", "扫描、选择、连接、校验、分片写入、提交、读回。", "扫描、连接、写入、校验、成功、可重试失败。"),
            ("卡片详情", "范围、检测点、时间、次数、离线标志、策略与写卡版本。", "服务端状态与本地写入状态分开显示。"),
            ("二次确认", "显示检测点、方向、动作和倒计时；同意或拒绝。", "待确认、已同意、已拒绝、超时、离线。"),
            ("安全管理", "冻结、挂失、恢复、解除绑定和异常记录。", "本地提交中、服务端已确认、待同步。"),
            ("临时通行", "用户主动开启手机身份载体模式，显示倒计时与耗电。", "后置开发；需目标真机验证 SLE 外围设备能力。"),
        ],
        [1800, 3900, 3660],
        font_size=8.8,
    )

    add_heading(doc, "4.2 推荐代码分层", 2)
    add_code_block(doc, "entry/src/main/ets/\n├─ pages/       卡包、邀请码、详情、写卡、确认、临时通行\n├─ models/      Card、Permission、Confirmation、CardWriteProtocol\n├─ stores/      SessionStore、CardStore、ConfirmationStore、SleStore\n├─ services/    MobileApiClient、MobileWebSocketClient、CardSleClient\n│               CardProtocolCodec、LocalStorageService、NotificationService\n└─ security/    TokenStore、WriteTicketVerifier")
    add_para(doc, "本地普通元数据使用 Preferences/RDB 保存；访问令牌、设备绑定秘密和敏感材料使用系统安全存储。写卡密钥不得进入日志、通知或普通配置文件。")

    add_heading(doc, "5. 统一数据与状态模型", 1)
    add_table(
        doc,
        ["对象", "核心字段", "设计要求"],
        [
            ("MobileCard", "permissionId、organizationId、displayName、status、policyVersion", "展示字段与硬件匿名字段分离。"),
            ("确认策略", "adminConfirmRequired、userConfirmEnabled、finalConfirmRequired", "final = admin OR user；admin 字段只读。"),
            ("写卡状态", "writeState、cardAnonId、credentialVersion、generation、lastWriteAt", "邀请码绑定状态与实体卡写入状态分离。"),
            ("确认请求", "requestId、eventId、detectorId、action、direction、expiresAt", "以 requestId 幂等；过期后不能补发同意。"),
            ("同步游标", "serverRevision、lastSyncAt、pendingOperations", "断网时保留用户操作，恢复后有序补传。"),
        ],
        [1800, 4100, 3460],
        font_size=9,
    )
    add_code_block(doc, "finalConfirmRequired = adminConfirmRequired || userConfirmEnabled")
    add_para(doc, "邀请兑换成功表示许可已绑定到用户；实体卡写入失败只将 writeState 标为 WRITE_FAILED。邀请码不恢复可用，但用户可对同一许可继续重试写卡，避免出现“假成功”。")

    add_heading(doc, "6. 手机专用后端接口", 1)
    add_callout(doc, "边界要求", "手机不得直接使用管理端全局 /api/cards、/api/confirmations/pending 或管理员令牌。必须增加按用户 subject 过滤的移动端路由和 WebSocket 频道。", fill="FDECEC", accent=RED)
    add_table(
        doc,
        ["接口", "用途", "关键约束"],
        [
            ("POST /api/mobile/pair", "建立手机设备会话。", "返回设备绑定、短期访问令牌和刷新令牌。"),
            ("POST /api/mobile/invites/preview", "预览邀请码对应许可。", "不消耗邀请码，不返回凭证密钥。"),
            ("POST /api/mobile/invites/redeem", "邀请码绑定用户。", "一次性/次数限制、审计、幂等。"),
            ("GET /api/mobile/cards", "获取当前用户许可。", "只能返回 token subject 所属卡片。"),
            ("PATCH /api/mobile/cards/:id/preferences", "设置用户自选确认。", "不得修改 adminConfirmRequired。"),
            ("POST /api/mobile/cards/:id/actions", "冻结、挂失、恢复或解绑。", "写审计；硬件同步状态单独返回。"),
            ("POST /api/mobile/cards/:id/write-package", "取得78字节凭证和短期写卡票据。", "绑定用户、许可、目标卡、版本、摘要和有效期。"),
            ("POST /api/mobile/cards/:id/write-receipt", "上传 Card C 提交和读回结果。", "成功后才把实体卡写入状态设为 ACTIVE。"),
            ("GET /api/mobile/confirmations/pending", "恢复未处理确认。", "仅返回该用户的未过期请求。"),
            ("POST /api/mobile/confirmations/:id/decision", "同意或拒绝。", "requestId 幂等；重复提交返回原结果。"),
            ("WS /ws/mobile", "推送确认、卡片状态和撤销消息。", "服务端按 subject 过滤 topic。"),
        ],
        [3350, 2850, 3160],
        font_size=8.5,
    )

    add_heading(doc, "6.1 局域网接入", 2)
    add_code_block(doc, "STARFOLLOW_HOST=0.0.0.0\nSTARFOLLOW_API_TOKEN=<管理端专用随机令牌>")
    add_para(doc, "现有后端在非回环地址监听时已经强制要求 API Token，这是良好基础。但移动端不能把管理员 Token 固化进 HAP，应新增独立移动端认证中间件。开发阶段可在同一局域网使用测试身份；正式演示应采用 HTTPS/WSS 或应用层加密。")

    add_heading(doc, "7. Card C 手机写卡协议改进", 1)
    add_para(doc, "继续复用现有 Card Service：Service 0x2C00，INFO 0x2C01，COMMAND 0x2C02，RESPONSE 0x2C03，STATUS 0x2C04。底层保持 Protocol V2 二进制帧，不把 JSON 放入 SLE。")
    add_heading(doc, "7.1 写卡票据", 2)
    add_table(
        doc,
        ["字段", "用途"],
        [
            ("requestId", "一次写卡操作的幂等标识。"),
            ("targetCardId", "票据只能用于指定 Card C。"),
            ("permissionId / credentialVersion", "绑定许可和版本，防止降级。"),
            ("credentialHash", "绑定实际78字节凭证内容。"),
            ("expireAt / nonce", "限制票据时效并防止重放。"),
            ("authorizationTag", "由后端写卡授权密钥生成，Card C 独立验证。"),
        ],
        [2600, 6760],
    )
    add_para(doc, "建议新增 AB_ROLE_PROVISIONER 表示手机写卡端。Card C 只有在 SLE 连接已加密认证、票据合法且未过期、目标卡和凭证摘要匹配时才接受写事务；不能直接取消现有 COMMAND 访问限制。")

    add_heading(doc, "7.2 正式写卡流程", 2)
    for text in [
        "手机向后端申请写卡包，取得78字节凭证、writeTicket 和 requestId。",
        "手机扫描名称/服务 UUID 匹配的 Card C，连接后读取 CARD_INFO。",
        "手机校验协议版本、cardId、容量、持久化能力和固件版本。",
        "手机提交写卡票据，Card C 验证后建立限时事务。",
        "依次发送 CREDENTIAL_BEGIN、CREDENTIAL_CHUNK 和 CREDENTIAL_COMMIT。",
        "收到 CREDENTIAL_RESULT 后调用 CREDENTIAL_LIST 读回 permissionId、版本和 generation。",
        "手机把读回结果上传后端，后端确认实体卡状态为 ACTIVE。",
    ]:
        numbered_paragraph(doc, text, decimal_id)

    add_heading(doc, "8. 二次确认闭环", 1)
    add_code_block(doc, "Detector A 形成完整事件\n  → Detector B 判断 finalConfirmRequired\n  → B 经 USB 上报 CONFIRM_REQUEST\n  → 后端保存并通过 /ws/mobile 推送指定用户\n  → 手机显示检测点、方向、动作和倒计时\n  → 手机提交 approve / reject\n  → 后端经 USB 发送 CONFIRM_RESULT\n  → B 执行、拒绝或超时，并上报最终事件")
    add_table(
        doc,
        ["情况", "手机表现", "硬件结果"],
        [
            ("管理端强制确认", "显示“管理方要求”，不能关闭。", "未收到 APPROVED 不得执行。"),
            ("用户自选确认", "用户可开启或关闭。", "开启后与管理端强制确认同等处理。"),
            ("拒绝", "立即显示已拒绝。", "B 保持关闭并记录原因。"),
            ("超时", "倒计时结束自动失效。", "B 返回 TIMEOUT，不接受迟到同意。"),
            ("手机/后端离线", "显示链路异常或离线。", "需要确认的执行默认失败。"),
        ],
        [2200, 3300, 3860],
    )

    add_heading(doc, "9. 安全、隐私与可靠性", 1)
    for text in [
        "管理令牌与移动端令牌分离；移动端接口按用户 subject 强制过滤。",
        "正式凭证密钥不写入页面、日志、通知、SQLite 明文字段或普通 Preferences。",
        "SLE 写卡使用加密认证连接、短期票据、版本检查、nonce 和 requestId 防重放。",
        "卡片只保存匿名 ID、许可、范围、策略、版本和转译密钥，不保存姓名、手机号等隐私。",
        "写卡必须收到 COMMIT 回执并读回列表；权限通过、确认通过、执行成功仍是三个独立状态。",
        "WebSocket 断开后通过 pending 接口和同步游标恢复，不能只依赖一次实时推送。",
        "Card C 正式版本必须启用并真板验证 NV 双槽持久化；RAM 固件只能用于联调。",
    ]:
        numbered_paragraph(doc, text, bullet_id)

    add_heading(doc, "10. 分阶段实施计划", 1)
    add_table(
        doc,
        ["阶段", "工作内容", "可验收交付"],
        [
            ("M0 工程恢复", "取得完整手机仓库，修复本机签名、权限和构建配置。", "可安装 HAP；基础页面可启动。"),
            ("M1 UI与本地数据", "卡包、邀请码、详情、确认、写卡状态机；本地持久化。", "重启后状态保持，Mock流程完整。"),
            ("M2 手机后端", "移动端鉴权、用户隔离、领卡、卡片同步和安全操作。", "手机与本地后端真实 API 联调。"),
            ("M3 真实确认", "WebSocket推送、同意/拒绝/超时/重连。", "使用现有 A/B 跑通确认闭环，无需第三块板。"),
            ("M4 SLE写卡", "Card C票据授权、ArkTS SLE客户端、分片、读回、重试。", "第三块板真写卡和认证通过。"),
            ("M5 临时通行", "手机主动开启身份载体模式、倒计时和退出。", "目标真机支持时完成一次检测事件。"),
            ("M6 三端验收", "异常恢复、安全、掉电、离线和完整演示。", "测试记录、安装包和演示视频。"),
        ],
        [1500, 4820, 3040],
        font_size=8.7,
    )
    add_callout(doc, "优先级", "先做 M0—M3。它们可以直接复用现有 A/B 和后端完成真实验证；M4 需要第三块板及 Card C 固件联合改动；M5 依赖目标华为设备是否开放所需的 NearLink 外围设备能力。", fill=PALE_BLUE, accent=NAVY)

    add_heading(doc, "11. 验收清单", 1)
    checks = [
        "邀请码可预览并一次性绑定；重复使用有明确错误。",
        "手机只能看到自己的许可、卡片和确认请求。",
        "应用重启后卡片和同步状态保持，断网恢复后能够补传。",
        "管理端强制二次确认在手机上不可关闭。",
        "确认同意、拒绝、超时、离线均能传回 B 并形成最终事件。",
        "写卡失败不会显示实体卡已激活，同一许可可以安全重试。",
        "Card C 写入后读回 permissionId、credentialVersion 和 generation 一致。",
        "冻结、挂失、过期和撤销最终同步到手机、后端和检测端。",
        "应用、后端和日志中不出现正式密钥或可直接识别个人的卡片明文。",
        "Card C 复位后正式 NV 固件仍能保留凭证和使用次数。",
    ]
    check_rows = [("□", text) for text in checks]
    add_table(doc, ["检查", "验收条件"], check_rows, [720, 8640], font_size=9.3)

    add_heading(doc, "12. 需项目组确认的冻结项", 1)
    add_table(
        doc,
        ["编号", "建议冻结项", "建议结论"],
        [
            ("D1", "移动端总体架构", "采用网络控制面 + SLE写卡面的双通道方案。"),
            ("D2", "手机与检测端关系", "手机不直连 B；确认经后端转发，写卡直连 Card C。"),
            ("D3", "协议边界", "HTTP/WebSocket 使用 JSON；SLE/USB 使用 Protocol V2 二进制帧。"),
            ("D4", "移动端权限", "新增手机专用令牌与用户隔离，不复用管理员全局 Token。"),
            ("D5", "Card C 写卡", "增加短期签名/HMAC票据，禁止无授权开放远端 COMMAND。"),
            ("D6", "实施顺序", "先领卡、同步和真实确认，再做真写卡，临时通行最后。"),
        ],
        [900, 3300, 5160],
        font_size=9,
    )

    # Keep table rows expandable and metadata explicit.
    doc.core_properties.title = "SLE-ID 移动端设计方案"
    doc.core_properties.subject = "HarmonyOS 移动端与硬件端、管理端通信及实施设计"
    doc.core_properties.author = "SLE-ID 项目组"
    doc.core_properties.keywords = "SLE, NearLink, HarmonyOS, Card C, Detector A, Detector B"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()

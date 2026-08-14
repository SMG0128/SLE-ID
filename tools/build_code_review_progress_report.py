from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "SLE-ID代码复查与初版文档进度对比_2026-08-09.docx"

BLUE = "2E74B5"
DARK = "1F4D78"
LIGHT = "F2F4F7"
GREEN = "548235"
GREEN_LIGHT = "E2F0D9"
AMBER = "C65911"
AMBER_LIGHT = "FCE4D6"
RED = "C00000"
RED_LIGHT = "F4CCCC"
GRAY = "666666"
WHITE = "FFFFFF"
CONTENT_WIDTH = 9360


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic
    return run


def add_run(paragraph, text, **kwargs):
    return set_run_font(paragraph.add_run(text), **kwargs)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths):
    if sum(widths) != CONTENT_WIDTH:
        raise ValueError(f"table widths must total {CONTENT_WIDTH}, got {sum(widths)}")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        element = tbl_cell_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tbl_cell_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_numbering_definition(doc, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
        lvl.append(fmt)
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), "•" if bullet else f"%{level + 1}.")
        lvl.append(text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(720 + level * 360))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(720 + level * 360))
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        lvl.append(p_pr)
        if bullet:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), "Arial")
            r_fonts.set(qn("w:hAnsi"), "Arial")
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id, level=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def add_bullet(doc, text, num_id, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = doc.add_paragraph(style=style)
    add_run(paragraph, text)
    return paragraph


def add_number(doc, text, num_id, level=0):
    style = "List Number" if level == 0 else "List Number 2"
    paragraph = doc.add_paragraph(style=style)
    add_run(paragraph, text)
    return paragraph


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    style_settings = (
        ("Title", 24, DARK, 0, 8),
        ("Subtitle", 12, GRAY, 0, 12),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    )
    for name, size, color, before, after in style_settings:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name == "Title" or name.startswith("Heading")
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number", "List Number 2"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, "SLE-ID · 代码复查与进度审计", size=9, color="808080")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "内部项目文档  |  ", size=9, color="808080")
    page = OxmlElement("w:fldSimple")
    page.set(qn("w:instr"), "PAGE")
    footer._p.append(page)


def add_table(doc, headers, rows, widths, status_column=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_header(header)
    for index, label in enumerate(headers):
        set_cell_shading(header.cells[index], LIGHT)
        p = header.cells[index].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, label, bold=True, color=DARK, size=9.5)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            text = str(value)
            color = None
            bold = False
            if status_column == index:
                bold = True
                if any(x in text for x in ("通过", "完成", "就绪", "一致")) and "部分" not in text and "未" not in text:
                    set_cell_shading(cells[index], GREEN_LIGHT)
                    color = GREEN
                elif any(x in text for x in ("未完成", "阻塞", "高")):
                    set_cell_shading(cells[index], RED_LIGHT)
                    color = RED
                else:
                    set_cell_shading(cells[index], AMBER_LIGHT)
                    color = AMBER
            add_run(p, text, bold=bold, color=color, size=9.2)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_status_box(doc, label, text, color):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_shading(table.cell(0, 0), color)
    p0 = table.cell(0, 0).paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_after = Pt(0)
    add_run(p0, label, bold=True, color=WHITE, size=10.5)
    p1 = table.cell(0, 1).paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    add_run(p1, text, size=10.5)
    set_table_geometry(table, [1400, 7960])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_break(doc):
    doc.add_page_break()


def add_path_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    set_cell = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT)
    set_cell.append(shd)
    add_run(p, text, size=9, color=GRAY)
    return p


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    # The built-in Word list styles already use real numbering definitions and
    # remain more portable across Word/LibreOffice than dynamically injected IDs.
    bullet_id = None
    number_id = None
    props = doc.core_properties
    props.title = "SLE-ID 代码复查与初版文档进度对比"
    props.subject = "Detector A / Detector B / Card C / 管理前后端"
    props.author = "SLE-ID 项目组"
    props.keywords = "SLE, WS63, H3863, Card C, 代码审查, 进度对比"

    title = doc.add_paragraph(style="Title")
    add_run(title, "SLE-ID 代码复查与初版文档进度对比", bold=True, color=DARK, size=24)
    subtitle = doc.add_paragraph(style="Subtitle")
    add_run(subtitle, "Detector A · Detector B · Card C · 管理端前后端", size=12, color=GRAY)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    metadata = (
        ("审查日期", "2026-08-09"),
        ("基准文档", r"C:\Users\20741\Downloads\SLE无感项目三端开发分工与开发路线.docx"),
        ("硬件/后端工作区", str(ROOT)),
        ("管理端快照", r"D:\naiwa1\SLE-ID-Admin-Backend\SLE-ID-Admin-Backend"),
    )
    for row, pair in zip(meta.rows, metadata):
        set_cell_shading(row.cells[0], LIGHT)
        add_run(row.cells[0].paragraphs[0], pair[0], bold=True, color=DARK, size=9.5)
        add_run(row.cells[1].paragraphs[0], pair[1], size=9)
    set_table_geometry(meta, [1800, 7560])

    doc.add_heading("执行结论", level=1)
    add_status_box(doc, "代码健康", "本轮未发现会让现有主机回归链立即失败的新回归。硬件核心测试、后端测试、TypeScript 编译、前端生产构建和 C 静态分析均通过。", GREEN)
    add_status_box(doc, "项目阶段", "当前已经具备 A/B 双板演示、Card C 离板核心、单 USB 协议和本地后端骨架；但尚未达到初版文档的三端真实闭环与最终验收。", AMBER)
    add_status_box(doc, "主要阻塞", "真实 Channel Sounding、A 同时连接 B/Card、第三块 Card 实机、前端真实 API/WS、B 端持久化离线队列、外设反馈闭环。", RED)

    p = doc.add_paragraph()
    add_run(p, "一句话判断：", bold=True, color=DARK)
    add_run(p, "这是一个“核心链路健康、实物闭环未完成”的中期版本。适合继续联调和私人仓库协作，不适合按初版最终清单宣称整体验收完成。")
    add_path_note(doc, "本轮为复查与报告工作，没有修改业务源码。前端构建仅刷新了 D:\\naiwa1 快照中的 dist 产物。")

    add_page_break(doc)
    doc.add_heading("1. 复查范围与方法", level=1)
    add_bullet(doc, "逐项核查 A/B 公共协议、状态机、认证中继、B 权限权威、单 USB 网关、Card 存储/服务/认证及 WS63 平台适配。", bullet_id)
    add_bullet(doc, "核查管理后端的 SQLite、REST、WebSocket、串口协议、命令重试、确认事务、邀请码和安全边界。", bullet_id)
    add_bullet(doc, "核查 D:\\naiwa1 管理前端的 API 层、WebSocket 适配、页面覆盖和生产构建。", bullet_id)
    add_bullet(doc, "对照初版文档的 H0–H8、P0–P7、四条演示链和最终验收清单。", bullet_id)
    add_bullet(doc, "复核 outputs 中 5 个固件包和 1 个源码包的 SHA-256 记录。", bullet_id)

    doc.add_heading("2. 本轮验证结果", level=1)
    add_table(doc,
              ["对象", "执行内容", "结果", "说明"],
              [
                  ("硬件核心", "tests/run_tests.ps1", "通过", "A/B 协议与核心、Card 存储/服务/HMAC、三方互操作、A 中继、B 网关均通过"),
                  ("工具链", "Card 串口 4 类 DryRun、A/B 安全脚本语法、管理映射 SelfTest", "通过", "未使用实体第三板或真实串口"),
                  ("C 静态分析", "GCC 15.2 -fanalyzer，13 个核心 C 文件", "通过", "无编译器诊断；不覆盖 SDK 平台驱动运行时"),
                  ("管理后端", "TypeScript 编译 + 4 个测试文件", "9/9 通过", "API、数据库、硬件服务和协议测试通过"),
                  ("管理前端", "vue-tsc --noEmit + Vite 生产构建", "通过", "存在大于 500 kB 的 chunk 警告，不阻塞功能"),
                  ("代码一致性", "工作区 admin-server 与 D:\\naiwa1/server 源码、测试哈希比较", "一致", "两处后端无源码差异"),
                  ("发布产物", "5 个固件包 + 1 个源码包 SHA-256", "一致", "与 outputs/SHA256SUMS.txt 完全吻合"),
              ], [1450, 2800, 1200, 3910], status_column=2)
    p = doc.add_paragraph()
    add_run(p, "测试边界：", bold=True, color=AMBER)
    add_run(p, "本轮没有第三块实体板，未执行真实 Card SLE、真实 CS、真实 B 串口→后端→前端整链、掉电恢复、功耗和长时间运行测试。")

    add_page_break(doc)
    doc.add_heading("3. 代码问题清单", level=1)
    p = doc.add_paragraph()
    add_run(p, "结论：", bold=True)
    add_run(p, "没有发现新的 P0 级崩溃或越权执行回归；发现 3 个应在真实闭环前处理的 P1 问题，以及 2 个 P2 风险。")
    add_table(doc,
              ["编号", "级别", "问题与证据", "影响", "建议"],
              [
                  ("R-01", "P1", "确认先被置为 sending，成功/失败后才 finish；启动时没有恢复 sending。证据：admin-server/src/db.ts:578–595、runtime.ts:53–64。", "后端在发送确认期间异常退出，记录可能永久不再出现在 pending 列表。", "启动时把未过期 sending 恢复为 pending，过期项置 expired；新增进程重启回归测试。"),
                  ("R-02", "P1", "REST 与 /ws/events 没有认证中间件；操作者仅取自 x-operator 或请求体。证据：app.ts:10–18、routes/api.ts:8–18、ws/hub.ts:16–20。", "默认 127.0.0.1 时风险受限；若改为 0.0.0.0/局域网，任意可达客户端可操作许可、卡片、确认和备份。", "保持默认仅本机；若开放网络，先加会话/令牌、角色校验、WS 鉴权与 CSRF/Origin 策略。"),
                  ("R-03", "P1", "邀请码只实现列表、创建和撤销，没有 redeem/bind；max_uses 字段也没有消费路径。证据：db.ts:181–193、666–704，routes/api.ts:116–125。", "初版要求的一次性领卡、绑定人、使用次数和审计闭环无法完成。", "增加事务性核销接口，校验状态/有效期/次数，写 used_by、used_at、状态和审计。"),
                  ("R-04", "P2", "Card 先持久化 usage_count，再缓存并发送响应。证据：card_ws63/src/card_auth.c:195–224。", "若持久化成功后立刻掉电且响应未送达，缓存丢失；重试可能额外消耗一次次数。属于保守拒绝，不会导致越权，但会提前耗尽次数。", "正式次数受限场景采用持久化会话回执/计数预留，或明确接受 at-least-once 计数语义并增加掉电测试。"),
                  ("R-05", "P2", "前端构建成功但两个主 chunk 约 1.04 MB。", "Ark Web 首屏加载和低配置设备内存占用偏高。", "按路由拆包，Element Plus/图表库按需引入；不阻塞当前硬件联调。"),
              ], [650, 650, 3100, 2350, 2610])

    doc.add_heading("3.1 不是代码回归、但不能漏报的未完成项", level=2)
    add_bullet(doc, "A 端仍由 demo enter/reverse/timeout 注入观测，不是真实 Channel Sounding。证据：firmware/h3863/sle_ab/h3863_sle_ab.c:350–390。", bullet_id)
    add_bullet(doc, "A↔B 是真实 SLE；A↔Card 当前依赖串口桥，尚未实现 A 同时维持 B/Card 两条 SLE 链路。", bullet_id)
    add_bullet(doc, "B 的策略仅保存在 RAM，离线补传队列也是固定 RAM 队列；复位后会丢待上传记录和动态策略。", bullet_id)
    add_bullet(doc, "B 的执行器回调在 GPIO 拉高后立即返回 true，没有闸机/继电器反馈输入，因此只是动作请求成功，不是外设真实成功。证据：h3863_sle_ab.c:594–600。", bullet_id)
    add_bullet(doc, "管理前端全部业务 API 仍走 mock，WebSocket 仅为占位，且没有二次确认操作页面。", bullet_id)

    doc.add_heading("4. 初版硬件路线 H0–H8 对比", level=1)
    add_table(doc,
              ["阶段", "当前状态", "已具备", "仍缺少", "判定"],
              [
                  ("H0 基础通信", "部分完成", "A/B 实体 SLE、Protocol V2、角色/ID/ACK/重试；Card 独立 SLE 服务和串口桥核心", "三类设备同时通过 SLE 互认；A 双连接管理；第三板实机", "未达阶段交付"),
                  ("H1 卡片凭证", "核心就绪", "多许可、双槽 NV、版本/CRC、分片事务、摘要列表、冻结/撤销状态、三种固件配置", "手机 SLE 写卡/读回；第三板首次个性化与锁定包实测", "部分完成"),
                  ("H2 单 USB 网关", "代码就绪", "B 心跳、事件/告警/确认、策略、ACK/重试；后端串口驱动与 SQLite", "真实 B COM→后端→前端验收；复位不丢的 Flash 队列", "部分完成"),
                  ("H3 CS 探测", "未完成", "接口规划和模拟观测入口", "WS63 CS 原始数据、距离/信道特征、可重复测试数据", "阻塞 P4"),
                  ("H4 事件状态机", "部分完成", "进入/反向/超时/冷却等核心状态与单测，双板 demo 已验收", "真实 CS 驱动下的路过、停留、方向、重复目标和阈值标定", "未达实景交付"),
                  ("H5 身份安全", "核心较完整", "HMAC-SHA-256、随机挑战、计数、防重放、版本/有效期/状态、A 中继、B 权威", "真实三板 SLE；黑名单/撤销同步；生产密钥生命周期和掉电边界", "部分完成"),
                  ("H6 策略确认", "部分完成", "记录/执行/管理确认/用户确认标志/离线/拒绝报警基础子集", "范围、时间、次数、方向；手机确认链；五类结果整链实测", "未达阶段交付"),
                  ("H7 外设闭环", "演示级", "GPIO 高电平约 500 ms，重复事件不重复执行", "继电器/闸机接入、反馈输入、失败/超时判定与回执", "未完成闭环"),
                  ("H8 稳定低功耗", "部分完成", "重连、有限重试、去重、Card NV 构建、主机故障注入", "Card 功耗、并发目标、B Flash 补传、掉电恢复、寿命与长稳记录", "未达阶段交付"),
              ], [650, 1150, 2650, 3200, 1710], status_column=1)

    p = doc.add_paragraph()
    add_run(p, "硬件主线判断：", bold=True, color=DARK)
    add_run(p, "H1/H2/H5 的软件核心进展最靠前；H3 是当前最大技术阻塞，H0 的三节点真实 SLE 和 H7 的外设反馈是整机闭环阻塞。")

    doc.add_heading("5. 总体路线 P0–P7 对比", level=1)
    add_table(doc,
              ["阶段", "初版目标", "当前事实", "状态"],
              [
                  ("P0 架构冻结", "统一字段、消息、策略、确认优先级、单 USB", "Protocol V2 与基础映射已固定；完整 10 项策略和手机/Card 链仍需冻结", "基本完成"),
                  ("P1 三端 Mock", "手机/Web Mock，硬件最小通信", "Web Mock 完成；A/B 通信完成；Card 核心完成；手机端工程不在当前交付物中", "部分完成"),
                  ("P2 管理闭环", "数据库、发卡、邀请码、设备、日志、串口模拟", "SQLite/REST/WS/串口后端与测试完成；前端仍 Mock；邀请码无核销绑定", "部分完成"),
                  ("P3 写卡闭环", "邀请码领卡→SLE 写卡→回执→手机成功", "Card 分片写卡和串口工具完成；手机端、真实 SLE 写卡、邀请码绑定未完成", "未完成"),
                  ("P4 检测闭环", "Card→CS→事件→B→USB→管理日志", "A/B 后半链具备；Card-A 真链和 CS 前半链缺失，前端也未接真实日志", "未完成"),
                  ("P5 策略闭环", "记录/自动执行/强制确认/自选确认/拒绝报警", "协议与核心覆盖基础子集；手机确认、完整策略字段和真实外设反馈未通", "部分完成"),
                  ("P6 稳定演示", "断线、离线、挂失撤销、执行失败、重复触发、长稳", "A/B 重连/重复/超时已实机验证；掉电、Flash 补传、Card 实机、执行失败和长稳未完成", "部分完成"),
                  ("P7 二版准备", "服务器同步、多地点模型", "仅保留配置与接口思路；没有真实同步实现", "预留"),
              ], [1050, 2300, 4500, 1510], status_column=3)

    add_page_break(doc)
    doc.add_heading("5.1 四条演示链", level=2)
    add_table(doc,
              ["链路", "初版成功判定", "当前可证明范围", "结论"],
              [
                  ("A 普通签到", "授权经过，只记录不执行", "策略核心和模拟事件可实现；真实 Card+CS+前端日志未串通", "模拟可演示"),
                  ("B 自动执行", "授权经过，B 执行并收到成功回执", "A/B 实板 GPIO 动作通过；外设反馈和真实 Card+CS 缺失", "双板可演示"),
                  ("C 强制确认", "手机不可关闭；确认后执行，拒绝/超时失败", "B/后端确认协议具备；手机链、前端页面和崩溃恢复缺失", "部分可演示"),
                  ("D 未授权报警", "拒绝执行，管理端实时报警并记录", "B 拒绝/告警协议与后端入库具备；真实经过和前端实时链未验收", "部分可演示"),
              ], [700, 2600, 4350, 1710], status_column=3)

    doc.add_heading("6. 最终验收清单对比", level=1)
    add_table(doc,
              ["初版验收项", "状态", "当前证据", "差距"],
              [
                  ("三端统一 ID、状态和策略语义", "部分完成", "Protocol V2、event key、状态/原因枚举和后端映射", "手机端未核对；完整策略字段未进入硬件协议"),
                  ("一张卡多许可，按组织域/检测点匹配", "部分完成", "Card 多许可、organization/scope 字段和 B 认证请求", "真实三板匹配与管理发卡同步未验收"),
                  ("策略可组合，模板不硬编码业务", "部分完成", "管理数据模型可保存多项策略", "硬件仅执行基础 flags，scope/time/usage/direction 不支持"),
                  ("强制确认不可被用户关闭", "部分完成", "B 合并管理/用户确认标志", "手机端缺失；前端确认页缺失"),
                  ("有效经过后才记录/报警", "部分完成", "状态机核心和模拟测试通过", "真实 CS/传感器输入未接入"),
                  ("单 USB 正式链路稳定", "部分完成", "协议、解析、ACK/重试、后端串口驱动和自动测试", "真实 B COM 整链与长稳未验收；B 复位会丢 RAM 队列"),
                  ("过滤路过、停留、反向、重复", "部分完成", "状态机/双板 demo 覆盖反向、超时、重复和冷却", "真实信号阈值、多人/多目标场景未验收"),
                  ("权限、确认、执行三状态独立", "基本完成", "协议和日志分别保留 auth/confirm/execution", "外设真实反馈仍缺"),
                  ("挂失/冻结/过期/撤销/密钥/重放有明确处理", "核心完成", "Card/B 核心状态、原因码和异常测试", "管理状态到实体 Card/B 的同步未闭环"),
                  ("无公网服务器可完整演示", "部分完成", "本地后端/SQLite 可离线运行", "三端真实链、前端真实数据、掉电补传尚未完成"),
                  ("密钥/隐私不进入正式演示数据", "部分完成", "匿名 card ID；测试密钥位于 TEST_MODE 条件代码", "生产密钥注入、轮换和审计流程未建立"),
              ], [2550, 1250, 3130, 2430], status_column=1)

    doc.add_heading("7. 建议整改顺序", level=1)
    p = doc.add_paragraph()
    add_run(p, "优先原则：", bold=True, color=DARK)
    add_run(p, "先消除会让真实闭环失真的问题，再补能力宽度；第三板暂缺时，先完成不依赖第三板的后端和接口工作。")

    priorities = [
        "修复 R-01：后端启动恢复 sending 确认，并新增重启测试；同时明确 failed 是否允许人工重试。",
        "完成前端真实 REST/WS 适配与二次确认页面；用后端 MockSerial/协议测试跑通页面→后端→虚拟 B。",
        "实现邀请码 redeem/bind 事务和卡片/许可绑定；随后补管理 API 的本机安全策略或正式鉴权。",
        "在没有第三板期间，先设计并实现 A 的双 SLE 连接状态机接口、Card 会话路由和可替换 CS 采样抽象。",
        "第三板到位后依次验收：NV 个性化→锁定包→Card-A SLE→三方 HMAC→复位计数→真实 B USB 管理整链。",
        "最后推进真实 Channel Sounding、完整 10 项策略、B Flash 队列、外设反馈和 8–24 小时长稳/功耗测试。",
    ]
    for item in priorities:
        add_number(doc, item, number_id)

    doc.add_heading("7.1 下一次验收的最小通过条件", level=2)
    add_bullet(doc, "后端重启后，不存在永久 sending 的有效确认；过期确认不会重新执行。", bullet_id)
    add_bullet(doc, "前端不再读取 mock，事件、设备、许可、告警、串口状态和确认均来自真实 API/WS。", bullet_id)
    add_bullet(doc, "虚拟串口链可完整证明：B 心跳→事件→数据库→WebSocket→页面；页面确认→B COMMAND_RESULT。", bullet_id)
    add_bullet(doc, "所有未支持策略在 API 返回中明确标记，不允许界面显示“已下发成功”。", bullet_id)

    doc.add_heading("8. 产物与版本证据", level=1)
    add_table(doc,
              ["产物", "SHA-256"],
              [
                  ("detector_a_h3863_all.fwpkg", "0C80BFC5F86842624DA666EDE70204EEA486601F6FC8A6F0134350C553966291"),
                  ("detector_b_h3863_all.fwpkg", "DF58FA33A8F1DF826428F6A92508E84555E6C428D32BB89AD1FFF316C654F5F9"),
                  ("card_c_h3863_ram_all.fwpkg", "1ED0AED1CBF0F45468FA28D03C82B951031F1E20DE10C4C1048E71B942830B89"),
                  ("card_c_ws63_nv_provisioning_all.fwpkg", "9153D8374837ACEBAC9EA9DB79315583B16007F33672C3C2F02CA21408FFBE1F"),
                  ("card_c_ws63_nv_locked_load_only.fwpkg", "5421DAD234AF71293C3EBCCDBA5A2D8F5950B671081E15181551D83C95286113"),
                  ("SLE_AB_H3863_source.zip", "195C0890788F52E5FD5B937741D61A3E9A2AA3E6E3975FA2CA950294CFFAD476"),
              ], [3700, 5660])
    p = doc.add_paragraph()
    add_run(p, "说明：", bold=True)
    add_run(p, "上述哈希在本轮重新计算并与 outputs/SHA256SUMS.txt 核对一致。本轮没有重新执行官方 SDK 全量构建；固件可构建性依据 2026-08-09 已完成的官方 SDK 构建记录与未变化的发布产物。")

    add_page_break(doc)
    doc.add_heading("9. 最终结论", level=1)
    add_status_box(doc, "可继续开发", "现有代码基线健康，适合在当前版本上继续，不需要推倒重来。", GREEN)
    add_status_box(doc, "暂不整体验收", "在真实 CS、三节点 SLE、前端真实链和外设反馈完成前，不应按初版文档宣称 P4/P5/P6 或硬件 H0–H8 全部完成。", RED)
    add_status_box(doc, "最先处理", "后端确认恢复、前端真实 API/WS、邀请码核销、安全边界；这些工作不依赖第三块板，可以立即推进。", BLUE)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
